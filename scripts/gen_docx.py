#!/home/hermes/.venv/bin/python3
"""Generate a Word doc from /tmp/papers.json. Output: /tmp/ai-papers-summary.docx"""
import json, sys
from datetime import date
from pathlib import Path

INPUT = "/tmp/papers.json"
OUTPUT = "/tmp/ai-papers-summary.docx"

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
from docx.shared import Pt

papers = json.loads(Path(INPUT).read_text())
doc = Document()
today = date.today().strftime("%B %d, %Y")
doc.add_heading(f"Research Summary — {today}", level=0)
for p in papers:
    doc.add_heading(p["title"], level=2)
    if p.get("authors"):
        para = doc.add_paragraph()
        para.add_run(p["authors"]).italic = True
    doc.add_paragraph(p["summary"])
    doc.add_paragraph()
doc.save(OUTPUT)
print(f"DOCX written to {OUTPUT} ({len(papers)} papers)")
